"""
chunker.py - Smart document chunking strategies.
Supports Markdown heading-based, Python AST-based, and sliding-window chunking.
"""

import re
import ast
import hashlib
from typing import List, Dict, Optional
from pathlib import Path


class ChunkResult:
    """A single chunk extracted from a document."""

    def __init__(self, content: str, index: int, start_line: int = 0,
                 end_line: int = 0, chunk_type: str = "text",
                 heading: str = "", metadata: Dict = None):
        self.content = content
        self.index = index
        self.start_line = start_line
        self.end_line = end_line
        self.chunk_type = chunk_type
        self.heading = heading
        self.metadata = metadata or {}

    def to_dict(self) -> Dict:
        return {
            "content": self.content,
            "index": self.index,
            "start_line": self.start_line,
            "end_line": self.end_line,
            "chunk_type": self.chunk_type,
            "heading": self.heading,
            "metadata": self.metadata,
        }


def chunk_text(text: str, file_type: str, chunk_size: int = 500,
               overlap: int = 50) -> List[ChunkResult]:
    """Route to the appropriate chunking strategy based on file type."""
    if not text or not text.strip():
        return []

    file_type = file_type.lower().lstrip(".")

    if file_type in ("md", "markdown"):
        return chunk_markdown(text, chunk_size)
    elif file_type == "py":
        return chunk_python(text, chunk_size)
    elif file_type in ("js", "ts", "jsx", "tsx"):
        return chunk_sliding_window(text, chunk_size, overlap)
    elif file_type == "ipynb":
        return chunk_notebook(text, chunk_size)
    else:
        return chunk_sliding_window(text, chunk_size, overlap)


def chunk_markdown(text: str, max_chunk_size: int = 500) -> List[ChunkResult]:
    """Chunk Markdown by headings (H1-H3). Merge small sections."""
    lines = text.split("\n")
    sections: List[Dict] = []
    current_heading = ""
    current_lines: List[str] = []
    current_start = 1

    heading_re = re.compile(r"^(#{1,3})\s+(.+)$")

    for i, line in enumerate(lines):
        m = heading_re.match(line)
        if m and current_lines:
            sections.append({
                "heading": current_heading,
                "content": "\n".join(current_lines),
                "start_line": current_start,
                "end_line": i,
            })
            current_heading = m.group(2).strip()
            current_lines = [line]
            current_start = i + 1
        else:
            if m:
                current_heading = m.group(2).strip()
            current_lines.append(line)

    if current_lines:
        sections.append({
            "heading": current_heading,
            "content": "\n".join(current_lines),
            "start_line": current_start,
            "end_line": len(lines),
        })

    # Merge small sections, split large ones
    chunks: List[ChunkResult] = []
    buffer = ""
    buffer_heading = ""
    buffer_start = 1

    for sec in sections:
        if len(buffer) + len(sec["content"]) <= max_chunk_size:
            if not buffer:
                buffer_heading = sec["heading"]
                buffer_start = sec["start_line"]
            buffer += ("\n\n" if buffer else "") + sec["content"]
        else:
            if buffer.strip():
                chunks.append(ChunkResult(
                    content=buffer.strip(),
                    index=len(chunks),
                    start_line=buffer_start,
                    end_line=sec["start_line"] - 1,
                    chunk_type="markdown",
                    heading=buffer_heading,
                ))
            # If section itself is too large, split with sliding window
            if len(sec["content"]) > max_chunk_size:
                sub_chunks = chunk_sliding_window(
                    sec["content"], max_chunk_size, 50
                )
                for sc in sub_chunks:
                    sc.heading = sec["heading"]
                    sc.chunk_type = "markdown"
                    sc.start_line += sec["start_line"] - 1
                    sc.end_line += sec["start_line"] - 1
                    sc.index = len(chunks)
                    chunks.append(sc)
                buffer = ""
                buffer_heading = ""
            else:
                buffer = sec["content"]
                buffer_heading = sec["heading"]
                buffer_start = sec["start_line"]

    if buffer.strip():
        chunks.append(ChunkResult(
            content=buffer.strip(),
            index=len(chunks),
            start_line=buffer_start,
            end_line=sections[-1]["end_line"] if sections else 0,
            chunk_type="markdown",
            heading=buffer_heading,
        ))

    return chunks


def chunk_python(source: str, max_chunk_size: int = 500) -> List[ChunkResult]:
    """Chunk Python source by functions and classes using AST."""
    chunks: List[ChunkResult] = []

    try:
        tree = ast.parse(source)
    except SyntaxError:
        return chunk_sliding_window(source, max_chunk_size, 50)

    lines = source.split("\n")

    # Collect top-level and nested definitions
    nodes = []
    for node in ast.iter_child_nodes(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            nodes.append(node)

    if not nodes:
        return chunk_sliding_window(source, max_chunk_size, 50)

    # Add module-level code before first definition
    if nodes and nodes[0].lineno > 1:
        header = "\n".join(lines[:nodes[0].lineno - 1]).strip()
        if header:
            chunks.append(ChunkResult(
                content=header,
                index=0,
                start_line=1,
                end_line=nodes[0].lineno - 1,
                chunk_type="python_header",
                heading="module_header",
            ))

    for node in nodes:
        start = node.lineno
        end = node.end_lineno or start
        segment = "\n".join(lines[start - 1:end]).strip()

        if not segment:
            continue

        name = getattr(node, "name", "unknown")
        ntype = "function" if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) else "class"
        docstring = ast.get_docstring(node) or ""

        if len(segment) > max_chunk_size:
            sub_chunks = chunk_sliding_window(segment, max_chunk_size, 50)
            for sc in sub_chunks:
                sc.heading = f"{ntype}:{name}"
                sc.chunk_type = f"python_{ntype}"
                sc.start_line += start - 1
                sc.end_line += start - 1
                sc.index = len(chunks)
                sc.metadata = {"name": name, "type": ntype, "docstring": docstring[:200]}
                chunks.append(sc)
        else:
            chunks.append(ChunkResult(
                content=segment,
                index=len(chunks),
                start_line=start,
                end_line=end,
                chunk_type=f"python_{ntype}",
                heading=f"{ntype}:{name}",
                metadata={"name": name, "type": ntype, "docstring": docstring[:200]},
            ))

    return chunks


def chunk_notebook(text: str, max_chunk_size: int = 500) -> List[ChunkResult]:
    """Chunk Jupyter notebook JSON by cells."""
    import json as _json
    chunks: List[ChunkResult] = []

    try:
        nb = _json.loads(text)
    except _json.JSONDecodeError:
        return chunk_sliding_window(text, max_chunk_size, 50)

    cells = nb.get("cells", [])
    for i, cell in enumerate(cells):
        cell_type = cell.get("cell_type", "code")
        source_lines = cell.get("source", [])
        content = "".join(source_lines) if isinstance(source_lines, list) else source_lines

        if not content.strip():
            continue

        if len(content) > max_chunk_size:
            sub = chunk_sliding_window(content, max_chunk_size, 50)
            for sc in sub:
                sc.chunk_type = f"notebook_{cell_type}"
                sc.heading = f"cell_{i}"
                sc.index = len(chunks)
                chunks.append(sc)
        else:
            chunks.append(ChunkResult(
                content=content.strip(),
                index=len(chunks),
                chunk_type=f"notebook_{cell_type}",
                heading=f"cell_{i}",
                metadata={"cell_index": i, "cell_type": cell_type},
            ))

    return chunks


def chunk_sliding_window(text: str, chunk_size: int = 500,
                          overlap: int = 50) -> List[ChunkResult]:
    """General sliding-window chunking with overlap."""
    if not text.strip():
        return []

    chunks: List[ChunkResult] = []
    lines = text.split("\n")
    current_chunk: List[str] = []
    current_len = 0
    current_start = 1

    for i, line in enumerate(lines):
        line_len = len(line) + 1
        if current_len + line_len > chunk_size and current_chunk:
            chunk_text_str = "\n".join(current_chunk)
            chunks.append(ChunkResult(
                content=chunk_text_str.strip(),
                index=len(chunks),
                start_line=current_start,
                end_line=current_start + len(current_chunk) - 1,
                chunk_type="text",
            ))
            # Keep overlap lines
            overlap_chars = 0
            overlap_lines = []
            for ol in reversed(current_chunk):
                overlap_chars += len(ol) + 1
                overlap_lines.insert(0, ol)
                if overlap_chars >= overlap:
                    break
            current_chunk = overlap_lines
            current_len = sum(len(x) + 1 for x in current_chunk)
            current_start = i + 1 - len(current_chunk) + 1
        current_chunk.append(line)
        current_len += line_len

    if current_chunk:
        chunk_text_str = "\n".join(current_chunk)
        if chunk_text_str.strip():
            chunks.append(ChunkResult(
                content=chunk_text_str.strip(),
                index=len(chunks),
                start_line=current_start,
                end_line=current_start + len(current_chunk) - 1,
                chunk_type="text",
            ))

    return chunks


def compute_content_hash(text: str) -> str:
    """SHA256 hash of text content."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def extract_text_from_file(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    p = Path(file_path)
    suffix = p.suffix.lower()

    if suffix in (".md", ".txt", ".py", ".js", ".ts", ".jsx", ".tsx",
                   ".java", ".go", ".rs", ".c", ".cpp", ".h", ".hpp",
                   ".css", ".html", ".xml", ".json", ".yaml", ".yml",
                   ".toml", ".ini", ".cfg", ".sh", ".bash", ".zsh",
                   ".sql", ".r", ".rb", ".php", ".swift", ".kt"):
        return p.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".ipynb":
        return p.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        # Basic PDF text extraction (requires no extra deps)
        return _extract_pdf_text(file_path)
    elif suffix == ".docx":
        return _extract_docx_text(file_path)
    else:
        try:
            return p.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""


def _extract_docx_text(file_path: str) -> str:
    """Extract text from a .docx file using python-docx library."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs)
    except ImportError:
        return f"[DOCX file: {file_path} - install python-docx for full extraction]"
    except Exception as e:
        return f"[Error reading DOCX file: {file_path}: {str(e)}]"


def _extract_pdf_text(file_path: str) -> str:
    """Simple PDF text extraction without heavy dependencies."""
    try:
        import subprocess
        result = subprocess.run(
            ["pdftotext", file_path, "-"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: read raw and extract visible text
    try:
        with open(file_path, "rb") as f:
            content = f.read()
        # Very basic extraction of text between parentheses in PDF
        import re
        texts = re.findall(rb"\(([^)]+)\)", content)
        return " ".join(t.decode("latin-1", errors="replace") for t in texts[:500])
    except Exception:
        return f"[PDF file: {file_path} - install pdftotext for full extraction]"
